"""
Copyright 2024 Amazon.com, Inc. or its affiliates. All Rights Reserved.

This is AWS Content subject to the terms of the Customer Agreement
----------------------------------------------------------------------
File content:
    Streamlit application for document upload and processing using GenAI
    with Cognito authentication
"""

from __future__ import annotations

import json
import os
import tempfile
from typing import Any, Callable, Dict, List
import datetime
import io
import base64
import hmac
import hashlib

import boto3
import aiohttp
import requests
import streamlit as st
import uuid
import time
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd

# constants
API_URI = os.environ.get("API_URI")
WS_API_URI = os.environ.get("WS_API_URI")
DEFAULT_NEGATIVE_ANSWER_QUESTION = "Could not answer based on the provided documents. Please rephrase your question, reduce the relevance threshold, or ask another question."  # noqa: E501
DEFAULT_NEGATIVE_ANSWER_SUMMARY = "Could not summarize the document."  # noqa: E501
WS_SSL = (os.environ.get("WS_SSL", "True")) == "True"

# Cognito configuration
COGNITO_REGION = os.environ.get("COGNITO_REGION", "us-west-2")
COGNITO_USER_POOL_ID = os.environ.get("COGNITO_USER_POOL_ID", "us-west-2_9CMo3zaE9")  # Replace with your User Pool ID
COGNITO_APP_CLIENT_ID = os.environ.get("COGNITO_APP_CLIENT_ID", "5cpq6kgsajjojk22hrruaelaia")  # Replace with your App Client ID
#COGNITO_APP_CLIENT_SECRET = os.environ.get("COGNITO_APP_CLIENT_SECRET", "")  # Replace with your App Client Secret if you have one

# Initialize AWS clients
bedrock_client = boto3.client(service_name="bedrock-runtime", region_name="us-west-2")
dynamodb_client = boto3.client('dynamodb', region_name="us-west-2")
cognito_idp = boto3.client('cognito-idp', region_name=COGNITO_REGION)
DOCUMENTS_TABLE = "rfi-document-processor"  # Name of your DynamoDB table


#def get_secret_hash(username):
 #   """
 #   Generate a secret hash for Cognito authentication
 #   """
 #   if not COGNITO_APP_CLIENT_SECRET:
 #       return None
        
 #   msg = username + COGNITO_APP_CLIENT_ID
 #   dig = hmac.new(
 #       bytes(COGNITO_APP_CLIENT_SECRET, 'utf-8'),
 #       msg=bytes(msg, 'utf-8'),
 #       digestmod=hashlib.sha256
 #   ).digest()

 #   return base64.b64encode(dig).decode()



def authenticate_user(username, password):
    """
    Authenticate a user with Cognito
    """
    try:
        auth_params = {
            'USERNAME': username,
            'PASSWORD': password
        }
            
        response = cognito_idp.initiate_auth(
            AuthFlow='USER_PASSWORD_AUTH',
            ClientId=COGNITO_APP_CLIENT_ID,
            AuthParameters=auth_params
        )
        
        # Check if user needs to set a new password
        if 'ChallengeName' in response and response['ChallengeName'] == 'NEW_PASSWORD_REQUIRED':
            return {
                'success': False, 
                'challenge': 'NEW_PASSWORD_REQUIRED',
                'session': response['Session'],
                'message': 'You need to set a permanent password. Please enter a new password.'
            }
        
        # Successful authentication
        if 'AuthenticationResult' in response:
            return {
                'success': True,
                'id_token': response['AuthenticationResult']['IdToken'],
                'access_token': response['AuthenticationResult']['AccessToken'],
                'refresh_token': response['AuthenticationResult']['RefreshToken'],
                'expires_in': response['AuthenticationResult']['ExpiresIn']
            }
        else:
            return {'success': False, 'message': 'Authentication failed - no result returned'}
        
    except cognito_idp.exceptions.NotAuthorizedException:
        return {'success': False, 'message': 'Incorrect username or password'}
        
    except cognito_idp.exceptions.UserNotConfirmedException:
        return {'success': False, 'message': 'User is not confirmed'}
        
    except Exception as e:
        return {'success': False, 'message': str(e)}


def set_permanent_password(username, new_password, session):
    """
    Set a permanent password for a user who has a temporary password
    """
    try:
        response = cognito_idp.respond_to_auth_challenge(
            ClientId=COGNITO_APP_CLIENT_ID,
            ChallengeName='NEW_PASSWORD_REQUIRED',
            Session=session,
            ChallengeResponses={
                'USERNAME': username,
                'NEW_PASSWORD': new_password
            }
        )
        
        if 'AuthenticationResult' in response:
            return {
                'success': True,
                'id_token': response['AuthenticationResult']['IdToken'],
                'access_token': response['AuthenticationResult']['AccessToken'],
                'refresh_token': response['AuthenticationResult']['RefreshToken'],
                'expires_in': response['AuthenticationResult']['ExpiresIn']
            }
        else:
            return {'success': False, 'message': 'Failed to set permanent password'}
            
    except Exception as e:
        return {'success': False, 'message': f'Error setting permanent password: {str(e)}'}


def get_user_info(access_token):
    """
    Get user information using the access token
    """
    try:
        response = cognito_idp.get_user(
            AccessToken=access_token
        )
        
        user_attrs = {}
        for attr in response['UserAttributes']:
            user_attrs[attr['Name']] = attr['Value']
        
        return {
            'success': True,
            'username': response['Username'],
            'attributes': user_attrs
        }
        
    except Exception as e:
        return {'success': False, 'message': str(e)}


def verify_token(token):
    """
    Verify if a token is valid and not expired
    """
    try:
        response = cognito_idp.get_user(
            AccessToken=token
        )
        return True
    except:
        return False


def update_document_status(document_id, file_name, s3_url, status, metadata=None):
    """Update the document status in DynamoDB table"""
    timestamp = datetime.datetime.now().isoformat()
    
    item = {
        'document_id': {'S': document_id},
        'file_name': {'S': file_name},
        's3_url': {'S': s3_url},
        'status': {'S': status},
        'timestamp': {'S': timestamp}
    }
    
    # Add user information if authenticated
    if 'user_info' in st.session_state and st.session_state.user_info:
        item['username'] = {'S': st.session_state.user_info.get('username', 'unknown')}
    
    if metadata:
        item['metadata'] = {'S': json.dumps(metadata)}
    
    try:
        dynamodb_client.put_item(
            TableName=DOCUMENTS_TABLE,
            Item=item
        )
        return True
    except Exception as e:
        st.error(f"Failed to update document status in DynamoDB: {e}")
        return False


def get_document_history():
    """Retrieve all documents from DynamoDB table"""
    try:
        # If user is authenticated and not an admin, only show their documents
        if 'user_info' in st.session_state and st.session_state.user_info:
            username = st.session_state.user_info.get('username')
            user_attributes = st.session_state.user_info.get('attributes', {})
            is_admin = user_attributes.get('custom:role') == 'admin' if 'custom:role' in user_attributes else False
            
            if not is_admin:
                # Filter by username
                response = dynamodb_client.scan(
                    TableName=DOCUMENTS_TABLE,
                    FilterExpression='username = :username',
                    ExpressionAttributeValues={':username': {'S': username}}
                )
            else:
                # Admin sees all documents
                response = dynamodb_client.scan(TableName=DOCUMENTS_TABLE)
        else:
            # Fallback if user info not available
            response = dynamodb_client.scan(TableName=DOCUMENTS_TABLE)
        
        items = response.get('Items', [])
        
        # Convert DynamoDB items to a list of dictionaries
        documents = []
        for item in items:
            doc = {
                'file_name': item.get('file_name', {}).get('S', ''),
                's3_url': item.get('s3_url', {}).get('S', ''),
                'status': item.get('status', {}).get('S', ''),
                'timestamp': item.get('timestamp', {}).get('S', ''),
                'document_id': item.get('document_id', {}).get('S', ''),
                'username': item.get('username', {}).get('S', 'unknown')
            }
            
            # Parse metadata if available
            if 'metadata' in item:
                try:
                    metadata = json.loads(item['metadata']['S'])
                    doc['metadata'] = metadata
                except (json.JSONDecodeError, KeyError):
                    doc['metadata'] = {}
            
            documents.append(doc)
        
        # Sort documents by timestamp (newest first)
        documents.sort(key=lambda x: x['timestamp'], reverse=True)
        return documents
    
    except Exception as e:
        st.error(f"Failed to retrieve document history: {e}")
        return []


def get_questions(document_chunk):
    from botocore.exceptions import ClientError
    
    prompt = f""" You are business person working on request of information requests.
    Your Job is to extract questions that meet the criteria in <instructions></instructions>.
    

    <document>
    {document_chunk}
    </document>             


    <instructions>
    The document contains multiple sections.
    The document contains a list of questions and subquestions that are marked with numbers and letters.
    Ignore the page number infront of the questions.
    Ignore Guidance.
    Ignore penalties.
    Ignore Warnings.
    Ignore Non-compliance with this notice. 
    Ignore Intentional obstruction or delay.
    Ignore instructions.
    Ignore Definitions and interpretations.
    </instructions>
    """
    
    # Invoke the model with the prompt
    model_id = "anthropic.claude-3-sonnet-20240229-v1:0"
    #model_id = "anthropic.claude-3-5-sonnet-20241022-v2:0"
    request_body = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 3200,
        "temperature": 0,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt,
                    },
                ],
            }
        ],
    }

    try:
        response = bedrock_client.invoke_model(
            modelId=model_id,
            body=json.dumps(request_body),
        )

        # Process the response
        result = json.loads(response.get("body").read())
        output_list = result.get("content", [])
        
        for output in output_list:
            if not output["text"].startswith("Unfortunately"):
                return output["text"]
            else: 
                return ""
                
    except ClientError as err:
        st.error(f"Error invoking Claude model: {str(err)}")
        raise


def get_answer(questions):
    #st.info("Processing questions to find answers...")
    client = boto3.client("bedrock-agent-runtime", region_name="us-west-2")

    #model_id = "anthropic.claude-3-haiku-20240307-v1:0"
    model_arn = 'arn:aws:bedrock:us-west-2::foundation-model/anthropic.claude-3-haiku-20240307-v1:0'
    #model_arn = 'arn:aws:bedrock:us-west-2::foundation-model/amazon.nova-pro-v1:0'
    kb_id = os.environ.get("KNOWLEDGE_BASE_ID", "ZINMEUUJPQ")  # Fallback to hardcoded for existing deployments
    numberOfResults = 5
    
    # Fixed prompt template with literal placeholders
    promptTemplate = """ Human: You are business person working on answering request of information requests. 
    Your Job is to extract answers that meet the criteria in <instructions></instructions> from the $search_results$ which contains numbered answers to questions. Extract the responses, keeping the numbers intact, and put each numbered response on a new line.

    Here are the search results in numbered order:
    $search_results$

    $output_format_instructions$

    <instructions> 
    Look for the exact match to that question or subquestions and provide all the answers and using the existing content as is and keep original text intact 
    When you find the exact match then the answeres will start with same number of the questions.
    Ensure to Provide all the answers that starts with same question number.
    Ignore the references in the end of the page which alwayes start with number then See the the url.
    Dont include the question as part of the result. 
    Include the answer number.
    </instructions>
    
    """

    llm_generated_questions = questions
    llm_generated_questions_array = llm_generated_questions.split("\n\n")
    answers = []
    
    #progress_bar = st.progress(0)
    total_questions = len(llm_generated_questions_array[1:])
    
    for i, llm_generated_question in enumerate(llm_generated_questions_array[1:]):
        #progress_bar.progress((i + 1) / total_questions)
        answer = {"question": "", "answer": ""}
        try:
            response = client.retrieve_and_generate(
                input={
                    'text': llm_generated_question
                },
                retrieveAndGenerateConfiguration={
                    'type': 'KNOWLEDGE_BASE',
                    'knowledgeBaseConfiguration': {
                        'knowledgeBaseId': kb_id,
                        'modelArn': model_arn,
                        'retrievalConfiguration': {
                            'vectorSearchConfiguration': {
                                'numberOfResults': numberOfResults,
                                'overrideSearchType': "SEMANTIC",
                            }
                        },
                        'generationConfiguration': {
                            'promptTemplate': {
                                'textPromptTemplate': promptTemplate
                            },
                            "inferenceConfig": { 
                               "textInferenceConfig": { 
                                  "maxTokens": 4000,
                                  "stopSequences": [ "Observation" ],
                                  "temperature": 0.0,
                                  "topP": 1
                               }
                            }
                        }
                    }
                },
            )
            
            print("response-retrieve_and_generate", response)
            response_answer = response["output"]["text"]
            answer["question"] = (llm_generated_question)
            
            if 'Sorry, I am unable' not in response_answer and "The search results do not contain" not in response_answer and "No answer" not in response_answer and "I could not find" not in response_answer and "there is no relevant information to extract for this question" not in response_answer: 
                answer["answer"] = "Answer: " + response_answer
                if len(response['citations'][0]['retrievedReferences']) > 0:
                    answer["metadata"] = "url: " + response['citations'][0]['retrievedReferences'][0]['location']['s3Location']['uri']
                else:
                    answer["metadata"] = " "
            else:
                answer["answer"] = " "
                answer["metadata"] = " "
                
        except Exception as e:
            st.error(f"Error during API call: {str(e)}")
            answer["answer"] = "Error: Unable to process the question."
            answer["metadata"] = " "
            
        answers.append(answer)
    
    return answers


def upload_to_s3(file):
    """Upload file to S3 bucket and return S3 URL"""
    s3 = boto3.client('s3')
    bucket_name = os.environ.get("DOCUMENT_BUCKET", "new-rfi-s3-bucket")  # Use environment variable
    
    # Create a unique file name
    #file_extension = file.name.split('.')[-1]
    #unique_filename = f"uploaded/{str(uuid.uuid4())}.{file_extension}"
    
    # Use the original file name but ensure it's safe for S3
    safe_filename = file.name.replace(" ", "_")  # Replace spaces with underscores

    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(file.getvalue())
        tmp_path = tmp.name
    
    try:
        #s3.upload_file(tmp_path, bucket_name, unique_filename)
        s3.upload_file(tmp_path, bucket_name, safe_filename)
        os.unlink(tmp_path)  # Remove temporary file
        #return f"s3://{bucket_name}/{unique_filename}"
        return f"s3://{bucket_name}/{safe_filename}"
    except Exception as e:
        st.error(f"Failed to upload to S3: {e}")
        os.unlink(tmp_path)
        return None


def run_chunk_document(process_message: Callable, docurl, document_id) -> None:
    try:
        from langchain_community.document_loaders import AmazonTextractPDFLoader
        from langchain_text_splitters import TokenTextSplitter
        
        with st.spinner("Processing document..."):
            # Update status to processing
            update_document_status(document_id, st.session_state.file_name, docurl, "PROCESSING")
            
            loader = AmazonTextractPDFLoader(docurl)
            documents = loader.load_and_split()

            text_splitter = TokenTextSplitter(chunk_size=1000, chunk_overlap=0)
            split_documents = text_splitter.split_documents(documents)
            
            st.info(f"Document processed into {len(split_documents)} text chunks")
            
            questionandanswer = []
            counter = 0
            max_chunks = len(split_documents)
            
            progress_bar = st.progress(0)
            
            for i, document_chunk in enumerate(split_documents[:max_chunks]):
                progress_bar.progress((i + 1) / max_chunks)
                output = get_questions(document_chunk)
                if output:  # Only process if questions were found
                    result = get_answer(output)
                    questionandanswer.extend(result)
                counter += 1
                if counter >= max_chunks:
                    break
            
            message = {
                "user_message": "",
                "bot_message": questionandanswer,
                "docs": questionandanswer
            }
            
            process_message(message)
            
            # Update status to completed
            metadata = {
                "chunks_processed": counter,
                "total_chunks": max_chunks,
                "answers_found": len(questionandanswer)
            }
            update_document_status(document_id, st.session_state.file_name, docurl, "COMPLETED", metadata)
            
    except Exception as e:
        st.error(f"An error occurred in run_chunk_document: {str(e)}")
        # Update status to error
        update_document_status(document_id, st.session_state.file_name, docurl, "ERROR", {"error_message": str(e)})


def process_message(message):
    st.session_state.processed_results = message["bot_message"]


def display_results():
    """Display the processed results in a structured format"""
    if "processed_results" in st.session_state and st.session_state.processed_results:
        st.subheader("Document Analysis Results")
        
        for i, result in enumerate(st.session_state.processed_results):
            with st.expander(f"Question {i+1}: {result['question'][:100]}..."):
                #st.markdown("### Question")
                #st.write(result["question"])
                
                st.markdown("### Answer")
                if result["answer"].strip() != "":
                    st.write(result["answer"])
                else:
                    st.write("No answer found for this question.")
                
                if "metadata" in result and result["metadata"].strip() != "":
                    st.markdown("### Source")
                    st.write(result["metadata"])


def create_word_document(results):
    """Create a Word document with the processed results"""
    from docx.shared import RGBColor
    
    doc = Document()
    
    # Add a title
    title = doc.add_heading('Document Analysis Results', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add timestamp
    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    timestamp_para = doc.add_paragraph(f'Generated on: {current_time}')
    timestamp_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Add space
    
    # Add results
    for i, result in enumerate(results):
        # Add question header with number
        question_header = doc.add_heading(f'Question {i+1}:', level=1)
        # Use RGBColor object instead of tuple
        for run in question_header.runs:
            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
        # Add question text
        question_para = doc.add_paragraph()
        question_run = question_para.add_run(result['question'])
        question_run.bold = True
        
        # Add answer header
        doc.add_heading('Answer:', level=2)
        
        # Add answer text
        if result["answer"].strip() != "":
            doc.add_paragraph(result["answer"])
        else:
            no_answer = doc.add_paragraph('No answer found for this question.')
            for run in no_answer.runs:
                run.italic = True
        
        # Add source/metadata if available
        if "metadata" in result and result["metadata"].strip() != "":
            doc.add_heading('Source:', level=2)
            doc.add_paragraph(result["metadata"])
        
        # Add separator between questions
        if i < len(results) - 1:
            doc.add_paragraph()
            separator = doc.add_paragraph('_' * 50)
            separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()
    
    # Return the document as bytes for download
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def display_login_page():
    """Display the login page"""
    st.title("RFI Document Processor")
    
    # Check if we're in password change mode
    if 'password_change_required' in st.session_state and st.session_state.password_change_required:
        st.subheader("Set Permanent Password")
        st.info("Your account has a temporary password. Please set a permanent password to continue.")
        
        col1, col2 = st.columns(2)
        with col1:
            new_password = st.text_input("New Password", type="password", help="Password must be at least 8 characters with uppercase, lowercase, numbers, and symbols")
            confirm_password = st.text_input("Confirm New Password", type="password")
            
            if st.button("Set Permanent Password"):
                if not new_password or not confirm_password:
                    st.error("Please enter and confirm your new password")
                elif new_password != confirm_password:
                    st.error("Passwords do not match")
                elif len(new_password) < 8:
                    st.error("Password must be at least 8 characters long")
                else:
                    with st.spinner("Setting permanent password..."):
                        result = set_permanent_password(
                            st.session_state.temp_username,
                            new_password,
                            st.session_state.temp_session
                        )
                        
                        if result['success']:
                            # Store tokens and set authenticated flag
                            st.session_state.authenticated = True
                            st.session_state.access_token = result['access_token']
                            st.session_state.id_token = result['id_token']
                            st.session_state.refresh_token = result['refresh_token']
                            
                            # Clean up temporary session data
                            del st.session_state.password_change_required
                            del st.session_state.temp_username
                            del st.session_state.temp_session
                            
                            # Get user info
                            user_info = get_user_info(result['access_token'])
                            if user_info['success']:
                                st.session_state.user_info = user_info
                                st.success(f"Password set successfully! Welcome, {user_info['username']}!")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.error(f"Error retrieving user information: {user_info['message']}")
                        else:
                            st.error(result['message'])
            
            if st.button("Cancel"):
                # Clear temporary session data
                if 'password_change_required' in st.session_state:
                    del st.session_state.password_change_required
                if 'temp_username' in st.session_state:
                    del st.session_state.temp_username
                if 'temp_session' in st.session_state:
                    del st.session_state.temp_session
                st.rerun()
        
        with col2:
            st.markdown("""
            ## Password Requirements
            
            Your new password must contain:
            - At least 8 characters
            - At least one uppercase letter
            - At least one lowercase letter
            - At least one number
            - At least one special character
            
            This ensures your account remains secure.
            """)
    else:
        st.subheader("Please log in to continue")
        
        col1, col2 = st.columns(2)
        with col1:
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            
            if st.button("Login"):
                if not username or not password:
                    st.error("Please enter both username and password")
                else:
                    with st.spinner("Authenticating..."):
                        auth_result = authenticate_user(username, password)
                        
                        if auth_result['success']:
                            # Store tokens and set authenticated flag
                            st.session_state.authenticated = True
                            st.session_state.access_token = auth_result['access_token']
                            st.session_state.id_token = auth_result['id_token']
                            st.session_state.refresh_token = auth_result['refresh_token']
                            
                            # Get user info
                            user_info = get_user_info(auth_result['access_token'])
                            if user_info['success']:
                                st.session_state.user_info = user_info
                                st.success(f"Welcome, {user_info['username']}!")
                                time.sleep(1)  # Give user time to see the message
                                st.rerun()
                            else:
                                st.error(f"Error retrieving user information: {user_info['message']}")
                        elif 'challenge' in auth_result and auth_result['challenge'] == 'NEW_PASSWORD_REQUIRED':
                            # Store temporary session data for password change
                            st.session_state.password_change_required = True
                            st.session_state.temp_username = username
                            st.session_state.temp_session = auth_result['session']
                            st.rerun()
                        else:
                            st.error(auth_result['message'])
    
    with col2:
        st.markdown("""
        ## Welcome to RFI Document Processor
        
        This application allows you to:
        - Upload and process PDF documents
        - Extract questions and find answers
        - View processing history
        - Download results as Word documents
        
        Please log in with your credentials to access the system.
        """)


def display_upload_tab():
    """Display the upload and processing tab"""
    st.header("Upload Document")
    uploaded_file = st.file_uploader("Upload a PDF document", type=["pdf"])
    
    col1, col2 = st.columns(2)
    
    with col1:
        if uploaded_file is not None:
            if st.button("Process Document"):
                with st.spinner("Uploading document to S3..."):
                    # Generate a document ID
                    document_id = str(uuid.uuid4())
                    st.session_state.document_id = document_id
                    st.session_state.file_name = uploaded_file.name
                    
                    # Upload to S3
                    s3_url = upload_to_s3(uploaded_file)
                    
                if s3_url:
                    st.session_state.uploaded_file_url = s3_url
                    st.success(f"File uploaded successfully to {s3_url}")
                    
                    # Record in DynamoDB that file was uploaded
                    update_document_status(document_id, uploaded_file.name, s3_url, "UPLOADED")
                    
                    # Process the document
                    run_chunk_document(process_message, s3_url, document_id)
                    st.success("Document processed successfully!")
                else:
                    st.error("Failed to upload the document.")
    
    with col2:
        if st.session_state.uploaded_file_url and st.button("Clear Results"):
            st.session_state.processed_results = []
            st.session_state.uploaded_file_url = None
            st.session_state.document_id = None
            st.session_state.file_name = None
            st.rerun()
    
    # Display results
    if st.session_state.processed_results:
        display_results()
        
        # Create Word document for download
        word_doc = create_word_document(st.session_state.processed_results)
        
        # Add download button for Word document
        if st.download_button(
            "Download Results as Word Document",
            data=word_doc,
            file_name="document_analysis_results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            st.success("Results downloaded successfully as Word document!")


def display_history_tab():
    """Display the document history tab"""
    st.header("Document Processing History")
    
    # Add a refresh button
    if st.button("Refresh History"):
        st.rerun()
    
    # Get document history from DynamoDB
    documents = get_document_history()
    
    if not documents:
        st.info("No document processing history found.")
        return
    
    # Create a DataFrame for the history table
    history_data = []
    for doc in documents:
        # Format timestamp for better readability
        try:
            timestamp = datetime.datetime.fromisoformat(doc['timestamp'])
            formatted_time = timestamp.strftime("%Y-%m-%d %H:%M:%S")
        except (ValueError, TypeError):
            formatted_time = doc['timestamp']
        
        # Extract additional metadata if available
        metadata_info = ""
        if 'metadata' in doc and isinstance(doc['metadata'], dict):
            if 'chunks_processed' in doc['metadata'] and 'total_chunks' in doc['metadata']:
                metadata_info = f"Processed {doc['metadata']['chunks_processed']}/{doc['metadata']['total_chunks']} chunks"
                if 'answers_found' in doc['metadata']:
                    metadata_info += f", {doc['metadata']['answers_found']} answers found"
            elif 'error_message' in doc['metadata']:
                metadata_info = f"Error: {doc['metadata']['error_message'][:50]}..."
        
        # Create a row for the table
        row = {
            "File Name": doc['file_name'],
            "Status": doc['status'],
            "Timestamp": formatted_time,
            "Processing Info": metadata_info,
            "User": doc.get('username', 'unknown')
        }
        history_data.append(row)
    
    # Create DataFrame and display as table
    history_df = pd.DataFrame(history_data)
    st.dataframe(history_df, use_container_width=True)
    
    # Add option to download history as CSV
    if st.download_button(
        "Download History as CSV",
        data=history_df.to_csv(index=False).encode('utf-8'),
        file_name="document_processing_history.csv",
        mime="text/csv"
    ):
        st.success("History downloaded successfully as CSV!")


def display_account_tab():
    """Display the user account tab"""
    st.header("My Account")
    
    if 'user_info' in st.session_state and st.session_state.user_info:
        user_info = st.session_state.user_info
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("User Information")
            st.write(f"**Username:** {user_info['username']}")
            
            # Display user attributes
            attributes = user_info.get('attributes', {})
            if 'email' in attributes:
                st.write(f"**Email:** {attributes['email']}")
            if 'name' in attributes:
                st.write(f"**Name:** {attributes['name']}")
            if 'custom:role' in attributes:
                st.write(f"**Role:** {attributes['custom:role']}")
            
            # Add logout button
            if st.button("Logout"):
                for key in ['authenticated', 'access_token', 'id_token', 'refresh_token', 'user_info']:
                    if key in st.session_state:
                        del st.session_state[key]
                st.session_state.current_page = "Upload"  # Reset to upload page after logout
                st.rerun()
        
        with col2:
            st.subheader("Session Information")
            st.write("Your session is currently active.")
            st.write("For security reasons, you will be automatically logged out after your session expires.")
    else:
        st.error("User information is not available. Please log out and log in again.")


# Main Streamlit App
def main():
    st.set_page_config(page_title="RFI Document Processor", layout="wide")
    
    # Initialize session state
    if "processed_results" not in st.session_state:
        st.session_state.processed_results = []
        
    if "uploaded_file_url" not in st.session_state:
        st.session_state.uploaded_file_url = None
        
    if "document_id" not in st.session_state:
        st.session_state.document_id = None
        
    if "file_name" not in st.session_state:
        st.session_state.file_name = None
        
    if "current_page" not in st.session_state:
        st.session_state.current_page = "Upload"
    
    # Check authentication
    if "authenticated" not in st.session_state or not st.session_state.authenticated:
        display_login_page()
        return
    
    # Verify token is still valid
    if "access_token" in st.session_state and not verify_token(st.session_state.access_token):
        st.error("Your session has expired. Please log in again.")
        for key in ['authenticated', 'access_token', 'id_token', 'refresh_token', 'user_info']:
            if key in st.session_state:
                del st.session_state[key]
        display_login_page()
        return
    
    # Sidebar navigation for authenticated users
    with st.sidebar:
        st.title("RFI Document Processor")
        
        # Show user info
        if 'user_info' in st.session_state and st.session_state.user_info:
            st.markdown(f"**Logged in as:** {st.session_state.user_info['username']}")
            
            # Check if user has admin role
            user_attributes = st.session_state.user_info.get('attributes', {})
            is_admin = user_attributes.get('custom:role') == 'admin' if 'custom:role' in user_attributes else False
            
            if is_admin:
                st.markdown("**Role:** Admin")
            else:
                st.markdown("**Role:** User")
        
        st.markdown("---")
        
        # Navigation options
        nav_options = ["Upload", "History", "Account"]
        selected = st.radio(
            "Navigation",
            options=nav_options,
            index=nav_options.index(st.session_state.current_page) if st.session_state.current_page in nav_options else 0
        )
        
        # Update current page in session state
        if selected != st.session_state.current_page:
            st.session_state.current_page = selected
            st.rerun()
        
        st.markdown("---")
        
        # Quick logout button
        if st.button("Logout", key="sidebar_logout"):
            for key in ['authenticated', 'access_token', 'id_token', 'refresh_token', 'user_info']:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()
        
        st.markdown("© 2024 Amazon.com, Inc.")
    
    # Main content based on selected page
    if st.session_state.current_page == "Upload":
        display_upload_tab()
    elif st.session_state.current_page == "History":
        display_history_tab()
    elif st.session_state.current_page == "Account":
        display_account_tab()


if __name__ == "__main__":
    main()